import json
import os
import sys
import tempfile
from pathlib import Path

import pytest

# Ensure backend is importable
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

# Use JSON backend for tests to avoid needing a database
os.environ["STORAGE_BACKEND"] = "json"
os.environ["OPENAI_API_KEY"] = "test"
os.environ["GEMINI_KEY"] = "test"
os.environ["DEVIN_KEY"] = "test"


@pytest.fixture(autouse=True)
def tmp_data_dir(tmp_path, monkeypatch):
    """Redirect JSON repo files to a temp directory for test isolation."""
    roles_file = tmp_path / "roles.json"
    users_file = tmp_path / "users.json"
    tenants_file = tmp_path / "tenants.json"
    workgroups_file = tmp_path / "workgroups.json"
    template_settings_file = tmp_path / "template_settings.json"
    workgroup_templates_file = tmp_path / "workgroup_templates.json"
    workgroup_users_file = tmp_path / "workgroup_users.json"
    workitems_file = tmp_path / "workitems.json"

    import repositories.json_repo as jr
    monkeypatch.setattr(jr, "ROLES_FILE", roles_file)
    monkeypatch.setattr(jr, "USERS_FILE", users_file)
    monkeypatch.setattr(jr, "TENANTS_FILE", tenants_file)
    monkeypatch.setattr(jr, "WORKGROUPS_FILE", workgroups_file)
    monkeypatch.setattr(jr, "TEMPLATE_SETTINGS_FILE", template_settings_file)
    monkeypatch.setattr(jr, "WORKGROUP_TEMPLATES_FILE", workgroup_templates_file)
    monkeypatch.setattr(jr, "WORKGROUP_USERS_FILE", workgroup_users_file)
    monkeypatch.setattr(jr, "WORKITEMS_FILE", workitems_file)
    monkeypatch.setattr(jr, "ANSWERSET_METADATA_FILE", tmp_path / "answerset_metadata.json")
    monkeypatch.setattr(jr, "AUDIT_LOG_FILE", tmp_path / "audit_log.json")

    import rate_limit
    rate_limit._attempts.clear()

    import auth_utils
    auth_utils._revoked_jtis.clear()

    return tmp_path


@pytest.fixture
def role_repo():
    from repositories.factory import get_role_repository
    return get_role_repository()


@pytest.fixture
def user_repo():
    from repositories.factory import get_user_repository
    return get_user_repository()


@pytest.fixture
def tenant_repo():
    from repositories.factory import get_tenant_repository
    return get_tenant_repository()


@pytest.fixture
def seeded_roles(role_repo):
    role_repo.create({"name": "admin", "description": "Administrator"})
    role_repo.create({"name": "staff", "description": "Staff member"})
    role_repo.create({"name": "approver", "description": "Approver"})
    role_repo.create({"name": "superadmin", "description": "Super administrator - manages tenants"})
    return role_repo


@pytest.fixture
def tenant_a(tenant_repo):
    return tenant_repo.create({
        "id": "tenant-a",
        "name": "Alpha Corp",
        "slug": "alpha",
        "active": "true",
        "created_at": "2026-01-01T00:00:00",
    })


@pytest.fixture
def tenant_b(tenant_repo):
    return tenant_repo.create({
        "id": "tenant-b",
        "name": "Beta Inc",
        "slug": "beta",
        "active": "true",
        "created_at": "2026-01-01T00:00:00",
    })


@pytest.fixture
def client(seeded_roles):
    from fastapi.testclient import TestClient
    from main import app
    with TestClient(app) as c:
        yield c


@pytest.fixture
def admin_token(client, user_repo, tenant_a):
    """Admin token for tenant A (alpha subdomain)."""
    from auth_utils import hash_password
    user_repo.create({
        "id": "test-admin-1",
        "username": "testadmin",
        "password": hash_password("pass1234"),
        "role": "admin",
        "name": "Test Admin",
        "tenant_id": "tenant-a",
    })
    resp = client.post(
        "/api/auth/login",
        data={"username": "testadmin", "password": "pass1234"},
        headers={"Host": "alpha.localhost:3000"},
    )
    return resp.json()["access_token"]


@pytest.fixture
def staff_token(client, user_repo, tenant_a):
    """Staff token for tenant A (alpha subdomain)."""
    from auth_utils import hash_password
    user_repo.create({
        "id": "test-staff-1",
        "username": "teststaff",
        "password": hash_password("pass1234"),
        "role": "staff",
        "name": "Test Staff",
        "tenant_id": "tenant-a",
    })
    resp = client.post(
        "/api/auth/login",
        data={"username": "teststaff", "password": "pass1234"},
        headers={"Host": "alpha.localhost:3000"},
    )
    return resp.json()["access_token"]


@pytest.fixture
def admin_token_tenant_b(client, user_repo, tenant_b):
    """Admin token for tenant B (beta subdomain)."""
    from auth_utils import hash_password
    user_repo.create({
        "id": "test-admin-b",
        "username": "testadmin",
        "password": hash_password("pass1234"),
        "role": "admin",
        "name": "Test Admin B",
        "tenant_id": "tenant-b",
    })
    resp = client.post(
        "/api/auth/login",
        data={"username": "testadmin", "password": "pass1234"},
        headers={"Host": "beta.localhost:3000"},
    )
    return resp.json()["access_token"]


@pytest.fixture
def superadmin_token(client, user_repo):
    """Superadmin token (admin subdomain, tenant_id=None)."""
    from auth_utils import hash_password
    user_repo.create({
        "id": "test-superadmin",
        "username": "testsuperadmin",
        "password": hash_password("pass1234"),
        "role": "superadmin",
        "name": "Super Admin",
        "tenant_id": None,
    })
    resp = client.post(
        "/api/auth/login",
        data={"username": "testsuperadmin", "password": "pass1234"},
        headers={"Host": "admin.localhost:3000"},
    )
    return resp.json()["access_token"]


def tenant_headers(slug, token):
    return {"Host": f"{slug}.localhost:3000", "Authorization": f"Bearer {token}"}


def admin_headers(token):
    return {"Host": "admin.localhost:3000", "Authorization": f"Bearer {token}"}


# ---------------------------------------------------------------------------
# Answerset / document helpers (shared by answerset, search and bulk tests)
# ---------------------------------------------------------------------------


@pytest.fixture
def answersets_env(tmp_path, monkeypatch):
    """Redirect answerset file storage (templates/submissions/generated docs)
    to a temp directory by patching ``BACKEND_ROOT`` in ``file_utils`` and the
    route modules that imported it by name. Returns the temp root.
    """
    root = tmp_path / "backend_root"
    root.mkdir()

    import file_utils
    import routes.answersets as answersets_route
    monkeypatch.setattr(file_utils, "BACKEND_ROOT", root)
    monkeypatch.setattr(answersets_route, "BACKEND_ROOT", root)

    return root


DEFAULT_COMPONENTS = [
    {"type": "string", "id": "customer_name", "label": "Customer name", "required": True},
    {"type": "string", "id": "work_description", "label": "Work done", "required": False, "multiline": True},
    {"type": "number", "id": "price", "label": "Price", "required": True, "min": 0},
]


def seed_template(root, tenant_id, template_id="tpl-1", name="Test Template", components=None):
    """Create a template (meta.json, interview.json, template.docx) on disk."""
    from docx import Document

    components = components if components is not None else DEFAULT_COMPONENTS
    tpl_dir = root / "data" / "templates" / tenant_id / template_id
    tpl_dir.mkdir(parents=True, exist_ok=True)

    (tpl_dir / "meta.json").write_text(json.dumps({
        "schemaVersion": 1,
        "id": template_id,
        "name": name,
        "active": True,
    }))
    (tpl_dir / "interview.json").write_text(json.dumps({
        "schemaVersion": 1,
        "id": f"{template_id}_interview",
        "version": 3,
        "components": components,
    }))

    doc = Document()
    doc.add_paragraph("Customer: {{customer_name}}")
    doc.add_paragraph("Work: {{work_description}}")
    doc.add_paragraph("Price: {{price}}")
    doc.save(str(tpl_dir / "template.docx"))

    return template_id


def valid_data(**overrides):
    data = {"customer_name": "Acme Ltd", "work_description": "Consulting", "price": 100}
    data.update(overrides)
    return data


def create_answerset(client, token, template_id="tpl-1", data=None, slug="alpha", **body):
    payload = {"template_id": template_id, "data": data if data is not None else valid_data(), **body}
    return client.post("/api/answersets/", json=payload, headers=tenant_headers(slug, token))
