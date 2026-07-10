"""Unit tests for document generation helpers in ``routes.answersets``.

These target the pure(-ish) building blocks behind the "generate a completed
document from an answerset" workflow — ``_generate_documents`` (docx render +
optional PDF conversion) and ``_calculate_completion`` (progress metric) —
without going through the HTTP layer. This is the core value of the product
(merging answers into a Word template) and was previously untested.
"""

import json
from pathlib import Path

import pytest
from docx import Document

import routes.answersets as answersets_route
from routes.answersets import _calculate_completion, _generate_documents


# ---------------------------------------------------------------------------
# _generate_documents
# ---------------------------------------------------------------------------


def _make_template_dir(tmp_path, body="Hello {{customer_name}}, total is {{price}}."):
    tpl_dir = tmp_path / "template"
    tpl_dir.mkdir()
    doc = Document()
    doc.add_paragraph(body)
    doc.save(str(tpl_dir / "template.docx"))
    return tpl_dir


class TestGenerateDocuments:
    def test_renders_docx_with_values(self, tmp_path):
        tpl_dir = _make_template_dir(tmp_path)
        generated = tmp_path / "generated"
        submission = {"id": "sub-1", "data": {"customer_name": "Acme", "price": 250}}

        docx_out, _pdf_out = _generate_documents({}, submission, tpl_dir, generated)

        assert docx_out == generated / "sub-1.docx"
        assert docx_out.exists()
        rendered = Document(str(docx_out))
        text = "\n".join(p.text for p in rendered.paragraphs)
        assert "Hello Acme, total is 250." in text

    def test_creates_generated_dir_if_missing(self, tmp_path):
        tpl_dir = _make_template_dir(tmp_path)
        generated = tmp_path / "does" / "not" / "exist"
        assert not generated.exists()
        docx_out, _ = _generate_documents(
            {}, {"id": "sub-2", "data": {"customer_name": "X", "price": 1}}, tpl_dir, generated
        )
        assert generated.exists()
        assert docx_out.exists()

    def test_missing_template_file_raises(self, tmp_path):
        empty_tpl_dir = tmp_path / "empty"
        empty_tpl_dir.mkdir()
        with pytest.raises(FileNotFoundError):
            _generate_documents(
                {}, {"id": "sub-3", "data": {}}, empty_tpl_dir, tmp_path / "generated"
            )

    def test_pdf_none_when_libreoffice_absent(self, tmp_path, monkeypatch):
        # Force the "no LibreOffice on PATH" branch regardless of the host.
        monkeypatch.setattr(answersets_route.shutil, "which", lambda _name: None)
        tpl_dir = _make_template_dir(tmp_path)
        _docx_out, pdf_out = _generate_documents(
            {}, {"id": "sub-4", "data": {"customer_name": "Y", "price": 2}}, tpl_dir, tmp_path / "gen"
        )
        assert pdf_out is None

    def test_pdf_produced_when_conversion_succeeds(self, tmp_path, monkeypatch):
        # Simulate LibreOffice being present and a successful conversion by
        # faking `which` and `subprocess.run` (returncode 0). The route treats a
        # zero return code as success and returns the expected pdf path.
        monkeypatch.setattr(answersets_route.shutil, "which", lambda _name: "/usr/bin/soffice")

        class _Result:
            returncode = 0

        monkeypatch.setattr(answersets_route.subprocess, "run", lambda *a, **k: _Result())
        tpl_dir = _make_template_dir(tmp_path)
        generated = tmp_path / "gen"
        _docx_out, pdf_out = _generate_documents(
            {}, {"id": "sub-5", "data": {"customer_name": "Z", "price": 3}}, tpl_dir, generated
        )
        assert pdf_out == generated / "sub-5.pdf"

    def test_pdf_none_when_conversion_fails(self, tmp_path, monkeypatch):
        monkeypatch.setattr(answersets_route.shutil, "which", lambda _name: "/usr/bin/soffice")

        class _Result:
            returncode = 1

        monkeypatch.setattr(answersets_route.subprocess, "run", lambda *a, **k: _Result())
        tpl_dir = _make_template_dir(tmp_path)
        _docx_out, pdf_out = _generate_documents(
            {}, {"id": "sub-6", "data": {"customer_name": "W", "price": 4}}, tpl_dir, tmp_path / "gen"
        )
        assert pdf_out is None


# ---------------------------------------------------------------------------
# _calculate_completion
# ---------------------------------------------------------------------------


class TestCalculateCompletion:
    FIELDS = [
        {"type": "string", "id": "a", "label": "A"},
        {"type": "string", "id": "b", "label": "B"},
        {"type": "number", "id": "c", "label": "C"},
        {"type": "number", "id": "d", "label": "D"},
    ]

    def test_no_fields_is_100(self):
        assert _calculate_completion({}, []) == 100.0

    def test_all_filled_is_100(self):
        data = {"a": "x", "b": "y", "c": 1, "d": 2}
        assert _calculate_completion(data, self.FIELDS) == 100.0

    def test_none_filled_is_0(self):
        assert _calculate_completion({}, self.FIELDS) == 0.0

    def test_partial_completion(self):
        data = {"a": "x", "b": "y"}  # 2 of 4
        assert _calculate_completion(data, self.FIELDS) == 50.0

    def test_empty_values_do_not_count(self):
        # Empty string, None, and empty list are treated as unfilled.
        data = {"a": "", "b": None, "c": [], "d": 5}
        assert _calculate_completion(data, self.FIELDS) == 25.0

    def test_dialog_fields_are_walked(self):
        fields = [
            {"type": "dialog", "id": "dlg", "title": "Section", "components": [
                {"type": "string", "id": "x", "label": "X"},
                {"type": "string", "id": "y", "label": "Y"},
            ]},
        ]
        assert _calculate_completion({"x": "filled"}, fields) == 50.0

    def test_repeat_fields_are_ignored(self):
        # Repeat groups are skipped by the completion walk, so only the plain
        # field counts toward the total.
        fields = [
            {"type": "string", "id": "a", "label": "A"},
            {"type": "repeat", "id": "rows", "label": "Rows", "components": [
                {"type": "string", "id": "line", "label": "Line"},
            ]},
        ]
        assert _calculate_completion({"a": "x"}, fields) == 100.0
