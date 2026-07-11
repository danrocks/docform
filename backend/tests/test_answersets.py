"""Tests for the answersets API (backend/routes/answersets.py).

Covers the answerset lifecycle end-to-end through the HTTP layer: creation
(with document generation), retrieval, listing/pagination, optimistic
concurrency control on update, cloning, sharing, deletion, download, audit
logging, and access-control / tenant-isolation rules.

The answerset routes read and write real files under ``BACKEND_ROOT`` (template
definitions, submission JSON, generated documents) in addition to the metadata
and audit repositories. The ``answersets_env`` fixture redirects all of that to
a temp directory so tests are fully isolated and never touch the repo tree.
"""

import json
from pathlib import Path

import pytest
from docx import Document

from tests.conftest import tenant_headers


# ---------------------------------------------------------------------------
# Fixtures / helpers
# ---------------------------------------------------------------------------


@pytest.fixture
def answersets_env(tmp_path, monkeypatch):
    """Redirect answerset file/repo storage to a temp directory.

    Patches:
      * ``BACKEND_ROOT`` in ``file_utils`` and ``routes.answersets`` so
        template/submission/generated files land under ``tmp_path``.
      * The answerset-metadata and audit-log JSON repo files.

    Returns the temp root used as ``BACKEND_ROOT``.
    """
    root = tmp_path / "backend_root"
    root.mkdir()

    import file_utils
    import routes.answersets as answersets_route
    monkeypatch.setattr(file_utils, "BACKEND_ROOT", root)
    monkeypatch.setattr(answersets_route, "BACKEND_ROOT", root)

    import repositories.json_repo as jr
    monkeypatch.setattr(jr, "ANSWERSET_METADATA_FILE", tmp_path / "answerset_metadata.json")
    monkeypatch.setattr(jr, "AUDIT_LOG_FILE", tmp_path / "audit_log.json")

    return root


DEFAULT_COMPONENTS = [
    {"type": "string", "id": "customer_name", "label": "Customer name", "required": True},
    {"type": "string", "id": "work_description", "label": "Work done", "required": False, "multiline": True},
    {"type": "number", "id": "price", "label": "Price", "required": True, "min": 0},
]


def seed_template(root, tenant_id, template_id="tpl-1", name="Test Template", components=None):
    """Create a template (meta.json, interview.json, template.docx) on disk."""
    components = components if components is not None else DEFAULT_COMPONENTS
    tpl_dir = root / "data" / "templates" / tenant_id / template_id
    tpl_dir.mkdir(parents=True, exist_ok=True)

    (tpl_dir / "meta.json").write_text(json.dumps({
        "schemaVersion": 1,
        "id": template_id,
        "name": name,
        "active": True,
    }))
    (tpl_dir / "interview.json").write_text(json.dumps({
        "schemaVersion": 1,
        "id": f"{template_id}_interview",
        "version": 3,
        "components": components,
    }))

    doc = Document()
    doc.add_paragraph("Customer: {{customer_name}}")
    doc.add_paragraph("Work: {{work_description}}")
    doc.add_paragraph("Price: {{price}}")
    doc.save(str(tpl_dir / "template.docx"))

    return template_id


def valid_data(**overrides):
    data = {"customer_name": "Acme Ltd", "work_description": "Consulting", "price": 100}
    data.update(overrides)
    return data


def create_answerset(client, token, template_id="tpl-1", data=None, **body):
    payload = {"template_id": template_id, "data": data if data is not None else valid_data(), **body}
    return client.post("/api/answersets/", json=payload, headers=tenant_headers("alpha", token))


# ---------------------------------------------------------------------------
# Creation
# ---------------------------------------------------------------------------


class TestCreateAnswerset:
    def test_admin_can_create_and_generates_docx(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        resp = create_answerset(client, admin_token)
        assert resp.status_code == 201
        body = resp.json()
        assert body["template_id"] == "tpl-1"
        assert body["template_name"] == "Test Template"
        assert body["status"] == "generated"
        assert body["submitted_by"] == "test-admin-1"
        # docx is generated even without LibreOffice; pdf may be None.
        assert body["docx_path"] and body["docx_path"].endswith(".docx")
        assert (answersets_env / body["docx_path"]).exists()
        # interviewVersion is captured from the interview file.
        assert body["interviewVersion"] == 3

    def test_created_docx_contains_rendered_values(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        body = create_answerset(client, admin_token, data=valid_data(customer_name="Zeta Corp")).json()
        rendered = Document(str(answersets_env / body["docx_path"]))
        text = "\n".join(p.text for p in rendered.paragraphs)
        assert "Zeta Corp" in text
        assert "Consulting" in text

    def test_staff_can_create(self, client, staff_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        resp = create_answerset(client, staff_token)
        assert resp.status_code == 201
        assert resp.json()["submitted_by"] == "test-staff-1"

    def test_create_persists_metadata(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]
        from repositories.factory import get_answerset_metadata_repository
        meta = get_answerset_metadata_repository().get_by_id(aid)
        assert meta is not None
        assert meta["version"] == 1
        assert meta["tenant_id"] == "tenant-a"
        assert meta["status"] == "generated"

    def test_create_missing_template_404(self, client, admin_token, answersets_env):
        # No template seeded.
        resp = create_answerset(client, admin_token, template_id="does-not-exist")
        assert resp.status_code == 404

    def test_create_invalid_data_400(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        # 'customer_name' is required; omit it.
        resp = create_answerset(client, admin_token, data={"price": 10})
        assert resp.status_code == 400

    def test_create_requires_auth(self, client, answersets_env):
        seed_template(answersets_env, "tenant-a")
        resp = client.post(
            "/api/answersets/",
            json={"template_id": "tpl-1", "data": valid_data()},
            headers={"Host": "alpha.localhost:3000"},
        )
        assert resp.status_code == 401


# ---------------------------------------------------------------------------
# Retrieval + listing/pagination
# ---------------------------------------------------------------------------


class TestGetAndList:
    def test_get_answerset_includes_completion_and_metadata(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]
        resp = client.get(f"/api/answersets/{aid}", headers=tenant_headers("alpha", admin_token))
        assert resp.status_code == 200
        body = resp.json()
        assert body["id"] == aid
        assert "completion_percentage" in body
        assert body["metadata"]["id"] == aid

    def test_get_missing_404(self, client, admin_token, answersets_env):
        resp = client.get("/api/answersets/nope", headers=tenant_headers("alpha", admin_token))
        assert resp.status_code == 404

    def test_list_pagination(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        for _ in range(5):
            assert create_answerset(client, admin_token).status_code == 201

        resp = client.get(
            "/api/answersets/?skip=0&limit=2",
            headers=tenant_headers("alpha", admin_token),
        )
        assert resp.status_code == 200
        body = resp.json()
        assert body["total"] == 5
        assert body["limit"] == 2
        assert body["skip"] == 0
        assert len(body["answersets"]) == 2

        resp2 = client.get(
            "/api/answersets/?skip=4&limit=2",
            headers=tenant_headers("alpha", admin_token),
        )
        assert len(resp2.json()["answersets"]) == 1

    def test_list_filter_by_template(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a", template_id="tpl-1")
        seed_template(answersets_env, "tenant-a", template_id="tpl-2", name="Other")
        create_answerset(client, admin_token, template_id="tpl-1")
        create_answerset(client, admin_token, template_id="tpl-2")

        resp = client.get(
            "/api/answersets/?template_id=tpl-2",
            headers=tenant_headers("alpha", admin_token),
        )
        assert resp.json()["total"] == 1
        assert resp.json()["answersets"][0]["template_id"] == "tpl-2"

    def test_limit_validation(self, client, admin_token, answersets_env):
        # limit above the allowed max (100) is rejected by FastAPI query validation.
        resp = client.get(
            "/api/answersets/?limit=500",
            headers=tenant_headers("alpha", admin_token),
        )
        assert resp.status_code == 422


# ---------------------------------------------------------------------------
# Update + optimistic concurrency
# ---------------------------------------------------------------------------


class TestUpdateConcurrency:
    def test_update_bumps_version(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]

        resp = client.put(
            f"/api/answersets/{aid}",
            json={"data": valid_data(customer_name="Updated Inc"), "version": 1},
            headers=tenant_headers("alpha", admin_token),
        )
        assert resp.status_code == 200
        assert resp.json()["metadata"]["version"] == 2
        assert resp.json()["data"]["customer_name"] == "Updated Inc"

    def test_stale_version_conflicts_409(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]

        # First update succeeds (version 1 -> 2).
        first = client.put(
            f"/api/answersets/{aid}",
            json={"data": valid_data(), "version": 1},
            headers=tenant_headers("alpha", admin_token),
        )
        assert first.status_code == 200

        # Second update with the now-stale version 1 must conflict.
        stale = client.put(
            f"/api/answersets/{aid}",
            json={"data": valid_data(), "version": 1},
            headers=tenant_headers("alpha", admin_token),
        )
        assert stale.status_code == 409

    def test_update_missing_404(self, client, admin_token, answersets_env):
        resp = client.put(
            "/api/answersets/nope",
            json={"data": valid_data(), "version": 1},
            headers=tenant_headers("alpha", admin_token),
        )
        assert resp.status_code == 404


# ---------------------------------------------------------------------------
# Clone / share / delete
# ---------------------------------------------------------------------------


class TestCloneShareDelete:
    def test_clone_creates_new_pending_answerset(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]

        resp = client.post(f"/api/answersets/{aid}/clone", headers=tenant_headers("alpha", admin_token))
        assert resp.status_code == 201
        clone = resp.json()
        assert clone["id"] != aid
        assert clone["status"] == "pending"
        assert clone["docx_path"] is None

        from repositories.factory import get_answerset_metadata_repository
        new_meta = get_answerset_metadata_repository().get_by_id(clone["id"])
        assert new_meta["cloned_from"] == aid
        assert new_meta["version"] == 1

    def test_owner_can_share(self, client, staff_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, staff_token).json()["id"]
        resp = client.put(
            f"/api/answersets/{aid}/share",
            json={"shared_with": ["user-x", "user-y"]},
            headers=tenant_headers("alpha", staff_token),
        )
        assert resp.status_code == 200
        assert resp.json()["shared_with"] == ["user-x", "user-y"]

    def test_non_owner_staff_cannot_share(self, client, staff_token, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        # Created by admin; staff (not owner, not admin) may not share.
        aid = create_answerset(client, admin_token).json()["id"]
        resp = client.put(
            f"/api/answersets/{aid}/share",
            json={"shared_with": ["user-x"]},
            headers=tenant_headers("alpha", staff_token),
        )
        assert resp.status_code == 403

    def test_owner_can_delete(self, client, staff_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, staff_token).json()["id"]
        resp = client.delete(f"/api/answersets/{aid}", headers=tenant_headers("alpha", staff_token))
        assert resp.status_code == 204
        assert client.get(
            f"/api/answersets/{aid}", headers=tenant_headers("alpha", staff_token)
        ).status_code == 404

    def test_non_owner_staff_cannot_delete(self, client, staff_token, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]
        resp = client.delete(f"/api/answersets/{aid}", headers=tenant_headers("alpha", staff_token))
        assert resp.status_code == 403


# ---------------------------------------------------------------------------
# Generate + download
# ---------------------------------------------------------------------------


class TestGenerateAndDownload:
    def test_regenerate_documents(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]
        resp = client.post(f"/api/answersets/{aid}/generate", headers=tenant_headers("alpha", admin_token))
        assert resp.status_code == 200
        assert resp.json()["status"] == "generated"

    def test_download_docx(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]
        resp = client.get(
            f"/api/answersets/{aid}/download/docx",
            headers=tenant_headers("alpha", admin_token),
        )
        assert resp.status_code == 200
        assert "wordprocessingml" in resp.headers["content-type"]
        assert len(resp.content) > 0

    def test_download_invalid_format_400(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]
        resp = client.get(
            f"/api/answersets/{aid}/download/txt",
            headers=tenant_headers("alpha", admin_token),
        )
        assert resp.status_code == 400

    def test_download_pdf_unavailable_404(self, client, admin_token, answersets_env):
        # Without LibreOffice the PDF is never produced, so download should 404.
        import shutil as _shutil
        if _shutil.which("libreoffice") or _shutil.which("soffice"):
            pytest.skip("LibreOffice present; PDF is generated")
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]
        resp = client.get(
            f"/api/answersets/{aid}/download/pdf",
            headers=tenant_headers("alpha", admin_token),
        )
        assert resp.status_code == 404


# ---------------------------------------------------------------------------
# Audit logging
# ---------------------------------------------------------------------------


class TestAuditLog:
    def test_audit_trail_records_operations(self, client, admin_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]
        # Trigger an access + a download to add more audit rows.
        client.get(f"/api/answersets/{aid}", headers=tenant_headers("alpha", admin_token))
        client.get(f"/api/answersets/{aid}/download/docx", headers=tenant_headers("alpha", admin_token))

        resp = client.get(f"/api/answersets/{aid}/audit", headers=tenant_headers("alpha", admin_token))
        assert resp.status_code == 200
        ops = {row["operation"] for row in resp.json()}
        assert {"create", "access", "download"} <= ops

    def test_staff_cannot_read_audit(self, client, staff_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, staff_token).json()["id"]
        resp = client.get(f"/api/answersets/{aid}/audit", headers=tenant_headers("alpha", staff_token))
        assert resp.status_code == 403


# ---------------------------------------------------------------------------
# Access control / tenant isolation
# ---------------------------------------------------------------------------


class TestAccessControl:
    def test_staff_cannot_access_others_answerset(self, client, admin_token, staff_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]
        # staff is neither owner, approver, nor shared-with.
        resp = client.get(f"/api/answersets/{aid}", headers=tenant_headers("alpha", staff_token))
        assert resp.status_code == 403

    def test_shared_user_can_access(self, client, admin_token, staff_token, answersets_env):
        seed_template(answersets_env, "tenant-a")
        aid = create_answerset(client, admin_token).json()["id"]
        # Share explicitly with the staff user id.
        client.put(
            f"/api/answersets/{aid}/share",
            json={"shared_with": ["test-staff-1"]},
            headers=tenant_headers("alpha", admin_token),
        )
        resp = client.get(f"/api/answersets/{aid}", headers=tenant_headers("alpha", staff_token))
        assert resp.status_code == 200

    def test_cross_tenant_answerset_not_visible(self, client, admin_token, admin_token_tenant_b, answersets_env):
        seed_template(answersets_env, "tenant-a")
        create_answerset(client, admin_token)
        # Tenant B admin lists their own (empty) answersets.
        resp = client.get("/api/answersets/", headers=tenant_headers("beta", admin_token_tenant_b))
        assert resp.status_code == 200
        assert resp.json()["total"] == 0
