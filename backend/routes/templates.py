import sys, os
import logging
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from fastapi import APIRouter, HTTPException, UploadFile, File, Depends, Form
from pydantic import BaseModel
from typing import Optional, List
import json, uuid, re, zipfile, base64
from pathlib import Path
from fastapi import HTTPException, APIRouter, Request
from fastapi.responses import FileResponse
from auth_utils import get_current_user, require_role
from question_schema import validate_questions
from config import settings
from AiResponseSaver import AiResponseSaver
from tenant_context import get_current_tenant, is_tenant_subdomain, verify_tenant_match
from file_utils import (
    BACKEND_ROOT,
    TEMPLATE_META_FILENAME,
    TEMPLATE_INTERVIEW_FILENAME,
    TEMPLATE_DOCX_FILENAME,
    get_tenant_data_dir,
    get_template_subdir,
    atomic_write_text,
    atomic_write_bytes,
    atomic_write_json,
    utcnow_iso,
    count_submissions,
)

from prompts.promptbuilder import _build_system_prompt
import providers  # noqa: F401 — triggers provider self-registration
from ai_providers import get_provider

logger = logging.getLogger(__name__)

SCHEMA_DIR = BACKEND_ROOT / "schema"

# Map provider name → schema file
PROVIDER_SCHEMA = {
    "devin":  "AiResponseSchemaFile.json",
    "gemini": "AiResponseSchema.json",
}

router = APIRouter()

OPENAI_SYSTEM_PROMPT = _build_system_prompt()


def get_templates_dir(request: Request) -> Path:
    return get_tenant_data_dir(request, "data", "templates")


def _get_submissions_dir_for_tenant(request: Request) -> Path:
    """Return the submissions dir for the current tenant (for counting)."""
    return get_tenant_data_dir(request, "data", "submissions")


def _load_template_with_interview(
    template_dir: Path,
    submissions_dir: Path | None = None,
) -> dict:
    """Load a template from its per-template subdirectory.

    Merges interview components/rules into the response and derives
    submissionCount from actual submission files when *submissions_dir*
    is provided.
    """
    meta_path = template_dir / TEMPLATE_META_FILENAME
    meta = json.loads(meta_path.read_text())

    interview_path = template_dir / TEMPLATE_INTERVIEW_FILENAME
    if interview_path.exists():
        interview = json.loads(interview_path.read_text())
        meta["fields"] = interview.get("components", [])
        meta["rules"] = interview.get("rules", [])
    else:
        logger.warning(
            "Interview file missing for template %s at %s",
            meta.get("id", "?"),
            interview_path,
        )
        meta["fields"] = []
        meta["rules"] = []

    # Derive submissionCount from actual files (#1)
    if submissions_dir is not None:
        meta["submissionCount"] = count_submissions(
            submissions_dir, meta.get("id", "")
        )

    return meta


def read_templates(templates_dir: Path, submissions_dir: Path | None = None) -> list:
    out = []
    for child in templates_dir.iterdir():
        if not child.is_dir():
            continue
        meta_path = child / TEMPLATE_META_FILENAME
        if not meta_path.exists():
            continue
        try:
            out.append(_load_template_with_interview(child, submissions_dir))
        except Exception:
            logger.exception("Failed to load template from %s", child)
    return sorted(out, key=lambda x: x.get("createdAt", x.get("created_at", "")), reverse=True)


def extract_placeholders_from_docx(path: Path) -> List[str]:
    """Extract {{placeholder}} tags from a docx file."""
    placeholders = set()
    try:
        with zipfile.ZipFile(path, "r") as z:
            for name in z.namelist():
                if name.endswith(".xml"):
                    text = z.read(name).decode("utf-8", errors="ignore")
                    found = re.findall(r'\{\{([^}]+)\}\}', text)
                    placeholders.update(f.strip() for f in found)
    except Exception:
        pass
    return sorted(placeholders)


def _validate_placeholders_vs_components(
    docx_path: Path,
    components: list,
) -> None:
    """Warn if docx placeholders and component IDs don't match (#7).

    Raises HTTPException(400) on mismatch so the user can fix it at
    upload time rather than discovering problems at render time.
    """
    placeholders = set(extract_placeholders_from_docx(docx_path))
    if not placeholders:
        return

    # Collect all component IDs (top-level only; repeat-group children
    # don't need top-level placeholders — the group id is the placeholder).
    component_ids = {c["id"] for c in components if "id" in c}

    missing_components = placeholders - component_ids
    missing_placeholders = component_ids - placeholders

    # Components inside repeat groups are allowed to lack top-level placeholders
    repeat_child_ids: set[str] = set()
    for c in components:
        if c.get("type") == "repeat":
            for child in c.get("components", []):
                if "id" in child:
                    repeat_child_ids.add(child["id"])
    missing_placeholders -= repeat_child_ids

    problems: list[str] = []
    if missing_components:
        problems.append(
            f"Document placeholders without matching interview components: "
            f"{', '.join(sorted(missing_components))}"
        )
    if missing_placeholders:
        problems.append(
            f"Interview components without matching document placeholders: "
            f"{', '.join(sorted(missing_placeholders))}"
        )
    if problems:
        raise HTTPException(status_code=400, detail="; ".join(problems))


@router.get("/ai-status")
def ai_status(current_user: dict = Depends(verify_tenant_match)):
    provider_name = os.environ.get("AI_PROVIDER", "devin")
    try:
        get_provider(provider_name, system_prompt=OPENAI_SYSTEM_PROMPT)
        return {"available": True}
    except Exception:
        return {"available": False}


@router.get("/")
def list_templates(request: Request, current_user: dict = Depends(verify_tenant_match)):
    templates_dir = get_templates_dir(request)
    submissions_dir = _get_submissions_dir_for_tenant(request)
    return read_templates(templates_dir, submissions_dir)


@router.get("/{template_id}")
def get_template(template_id: str, request: Request, current_user: dict = Depends(verify_tenant_match)):
    templates_dir = get_templates_dir(request)
    submissions_dir = _get_submissions_dir_for_tenant(request)
    template_dir = templates_dir / template_id
    meta_path = template_dir / TEMPLATE_META_FILENAME
    if not meta_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")
    return _load_template_with_interview(template_dir, submissions_dir)


@router.post("/")
async def create_template(
    request: Request,
    name: str = Form(...),
    description: str = Form(""),
    file: UploadFile = File(...),
    interview_json: Optional[str] = Form(None),
    current_user: dict = Depends(verify_tenant_match),
):
    if current_user["role"] not in ("admin", "superadmin"):
        raise HTTPException(status_code=403, detail="Not permitted")
    templates_dir = get_templates_dir(request)

    if not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    template_id = str(uuid.uuid4())
    template_dir = get_template_subdir(templates_dir, template_id)
    upload_path = template_dir / TEMPLATE_DOCX_FILENAME

    content = await file.read()
    atomic_write_bytes(upload_path, content)

    if interview_json:
        try:
            parsed = json.loads(interview_json)
        except json.JSONDecodeError as e:
            upload_path.unlink(missing_ok=True)
            raise HTTPException(status_code=400, detail=f"Invalid interview JSON: {e}")

        if isinstance(parsed, dict) and "components" in parsed:
            interview = parsed
        elif isinstance(parsed, list):
            interview = {
                "schemaVersion": 1,
                "id": f"{template_id}_interview",
                "version": 1,
                "components": parsed,
            }
        else:
            upload_path.unlink(missing_ok=True)
            raise HTTPException(status_code=400, detail="Invalid interview JSON: expected object with 'components' or a list of components")

        try:
            validate_questions(interview.get("components", []))
        except ValueError as e:
            upload_path.unlink(missing_ok=True)
            raise HTTPException(status_code=400, detail=str(e))

        _validate_placeholders_vs_components(upload_path, interview.get("components", []))
    else:
        placeholders = extract_placeholders_from_docx(upload_path)
        interview = {
            "schemaVersion": 1,
            "id": f"{template_id}_interview",
            "version": 1,
            "components": [
                {
                    "type": "string",
                    "id": p,
                    "label": p.replace("_", " ").title(),
                    "required": True,
                    "maxLength": 500,
                }
                for p in placeholders
            ],
        }

    # Remove title/description from interview — meta is the single source of truth (#6)
    interview.pop("title", None)
    interview.pop("description", None)
    interview.pop("$schema", None)

    atomic_write_json(template_dir / TEMPLATE_INTERVIEW_FILENAME, interview)

    meta = {
        "schemaVersion": 1,
        "id": template_id,
        "name": name,
        "description": description,
        "documentFile": TEMPLATE_DOCX_FILENAME,
        "interviewFile": TEMPLATE_INTERVIEW_FILENAME,
        "originalFilename": file.filename,
        "active": True,
        "createdAt": utcnow_iso(),
        "createdBy": current_user["id"],
        "updatedAt": None,
        "generationMethod": "upload",
    }

    atomic_write_json(template_dir / TEMPLATE_META_FILENAME, meta)

    submissions_dir = _get_submissions_dir_for_tenant(request)
    return _load_template_with_interview(template_dir, submissions_dir)


@router.put("/{template_id}")
def update_template(
    template_id: str,
    body: dict,
    request: Request,
    current_user: dict = Depends(verify_tenant_match),
):
    if current_user["role"] not in ("admin", "superadmin"):
        raise HTTPException(status_code=403, detail="Not permitted")
    templates_dir = get_templates_dir(request)
    template_dir = templates_dir / template_id
    meta_path = template_dir / TEMPLATE_META_FILENAME
    if not meta_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")
    meta = json.loads(meta_path.read_text())

    allowed_meta = {"name", "description", "active"}
    for k, v in body.items():
        if k in allowed_meta:
            meta[k] = v

    if "fields" in body:
        try:
            components = validate_questions(body["fields"])
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))

        interview_path = template_dir / TEMPLATE_INTERVIEW_FILENAME
        if interview_path.exists():
            interview = json.loads(interview_path.read_text())
        else:
            interview = {
                "schemaVersion": 1,
                "id": f"{template_id}_interview",
                "version": 1,
            }
        interview["components"] = components
        # Auto-increment interview version (#4)
        interview["version"] = interview.get("version", 0) + 1
        # Ensure no title/description in interview (#6)
        interview.pop("title", None)
        interview.pop("description", None)
        interview.pop("$schema", None)

        # Cross-validate placeholders vs components (#7)
        docx_path = template_dir / TEMPLATE_DOCX_FILENAME
        if docx_path.exists():
            _validate_placeholders_vs_components(docx_path, components)

        atomic_write_json(interview_path, interview)

    meta["updatedAt"] = utcnow_iso()
    atomic_write_json(meta_path, meta)

    submissions_dir = _get_submissions_dir_for_tenant(request)
    return _load_template_with_interview(template_dir, submissions_dir)


@router.delete("/{template_id}")
def delete_template(
    template_id: str,
    request: Request,
    current_user: dict = Depends(verify_tenant_match),
):
    if current_user["role"] not in ("admin", "superadmin"):
        raise HTTPException(status_code=403, detail="Not permitted")
    templates_dir = get_templates_dir(request)
    template_dir = templates_dir / template_id
    meta_path = template_dir / TEMPLATE_META_FILENAME
    if not meta_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")

    # Remove all files in the template subdirectory
    for f in template_dir.iterdir():
        f.unlink(missing_ok=True)
    template_dir.rmdir()

    return {"detail": "Deleted"}


class GenerateRequest(BaseModel):
    name: str
    description: Optional[str] = ""
    prompt: str


class RegenerateRequest(BaseModel):
    prompt: str


def _create_docx_from_content(document_content: str, output_path: Path):
    """Create a .docx file from AI-generated document content."""
    from docx import Document
    doc = Document()
    lines = document_content.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.isupper() and len(stripped) > 3 and not stripped.startswith("{{"):
            doc.add_heading(stripped, level=1)
        else:
            doc.add_paragraph(stripped)
    doc.save(str(output_path))


def _call_ai(prompt: str, model: str | None = None, tenant_id: str | None = None) -> dict:
    provider_name = os.environ.get("AI_PROVIDER", "devin")
    provider = get_provider(provider_name, system_prompt=OPENAI_SYSTEM_PROMPT)
    kwargs = {}
    if model:
        kwargs["model"] = model
    if tenant_id:
        kwargs["tenant_id"] = tenant_id
    return provider.call(prompt, **kwargs)

def _get_provider_format() -> str | None:
    """Read the document format from the provider's schema file."""
    provider_name = os.environ.get("AI_PROVIDER", "devin")
    schema_file = PROVIDER_SCHEMA.get(provider_name)
    if not schema_file:
        return None
    schema_path = SCHEMA_DIR / schema_file
    if not schema_path.exists():
        return None
    schema = json.loads(schema_path.read_text())
    return schema.get("properties", {}).get("document", {}).get("format")


@router.post("/generate")
def generate_template(
    request: Request,
    body: GenerateRequest,
    current_user: dict = Depends(verify_tenant_match),
):
    if current_user["role"] not in ("admin", "superadmin"):
        raise HTTPException(status_code=403, detail="Not permitted")
    templates_dir = get_templates_dir(request)

    tenant = get_current_tenant(request)
    ai_result = _call_ai(body.prompt, tenant_id=tenant["id"] if tenant else None)

    fmt = _get_provider_format()

    template_id = str(uuid.uuid4())
    template_dir = get_template_subdir(templates_dir, template_id)
    upload_path = template_dir / TEMPLATE_DOCX_FILENAME

    if fmt == "url":
        import httpx as _httpx

        doc_url = ai_result.get("document", "")
        int_url = ai_result.get("interview", "")
        if not doc_url or not int_url:
            raise HTTPException(status_code=500, detail="AI did not return document/interview URLs")

        def _normalize_devin_url(url: str) -> str:
            if "app.devin.ai/attachments/" in url:
                return url.replace("https://app.devin.ai/attachments/", "https://api.devin.ai/v1/attachments/")
            return url

        doc_url = _normalize_devin_url(doc_url)
        int_url = _normalize_devin_url(int_url)

        devin_key = os.environ.get("DEVIN_KEY", getattr(settings, "DEVIN_KEY", None))
        download_headers = {}
        if devin_key:
            download_headers["Authorization"] = f"Bearer {devin_key}"

        max_retries = 3

        doc_resp = None
        for attempt in range(max_retries):
            try:
                doc_resp = _httpx.get(doc_url, follow_redirects=True, timeout=60, headers=download_headers)
                doc_resp.raise_for_status()
                break
            except Exception as e:
                if attempt == max_retries - 1:
                    raise HTTPException(status_code=502, detail=f"Failed to download document from {doc_url} after {max_retries} attempts: {e}")
                import time
                time.sleep(2 ** attempt)

        atomic_write_bytes(upload_path, doc_resp.content)

        int_resp = None
        for attempt in range(max_retries):
            try:
                int_resp = _httpx.get(int_url, follow_redirects=True, timeout=60, headers=download_headers)
                int_resp.raise_for_status()
                break
            except Exception as e:
                if attempt == max_retries - 1:
                    raise HTTPException(status_code=502, detail=f"Failed to download interview from {int_url} after {max_retries} attempts: {e}")
                import time
                time.sleep(2 ** attempt)

        interview_data = int_resp.json()

        if isinstance(interview_data, dict):
            raw_questions = interview_data.get("components", interview_data.get("questions", []))
        elif isinstance(interview_data, list):
            raw_questions = interview_data
        else:
            raise HTTPException(status_code=500, detail="Unexpected interview format")

    elif fmt == "base64":
        doc_b64 = ai_result.get("document", "")
        int_b64 = ai_result.get("interview", "")
        if not doc_b64 or not int_b64:
            raise HTTPException(status_code=500, detail="AI did not return document/interview content")

        try:
            atomic_write_bytes(upload_path, base64.b64decode(doc_b64))
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to decode document: {e}")

        try:
            interview_text = base64.b64decode(int_b64).decode("utf-8")
            interview_data = json.loads(interview_text)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to decode interview: {e}")

        if isinstance(interview_data, dict):
            raw_questions = interview_data.get("components", interview_data.get("questions", []))
        elif isinstance(interview_data, list):
            raw_questions = interview_data
        else:
            raise HTTPException(status_code=500, detail="Unexpected interview format")

    else:
        document_content = ai_result.get("document_content", "")
        raw_questions = ai_result.get("questions", [])
        interview_data = None
        if not document_content:
            raise HTTPException(status_code=500, detail="AI did not generate document content")
        try:
            _create_docx_from_content(document_content, upload_path)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to create document: {e}")

    if not raw_questions:
        raise HTTPException(status_code=500, detail="AI did not generate interview questions")

    try:
        fields = validate_questions(raw_questions)
    except ValueError as e:
        raise HTTPException(status_code=500, detail=f"AI generated invalid questions: {e}")

    description = body.description or ai_result.get("summary", "")

    if isinstance(interview_data, dict) and "components" in interview_data:
        interview = interview_data
        interview["components"] = fields
    else:
        interview = {
            "schemaVersion": 1,
            "id": f"{template_id}_interview",
            "version": 1,
            "components": fields,
        }

    # Remove title/description from interview — meta is the single source of truth (#6)
    interview.pop("title", None)
    interview.pop("description", None)
    interview.pop("$schema", None)

    atomic_write_json(template_dir / TEMPLATE_INTERVIEW_FILENAME, interview)

    meta = {
        "schemaVersion": 1,
        "id": template_id,
        "name": body.name,
        "description": description,
        "documentFile": TEMPLATE_DOCX_FILENAME,
        "interviewFile": TEMPLATE_INTERVIEW_FILENAME,
        "originalFilename": f"{body.name.replace(' ', '_')}.docx",
        "active": True,
        "createdAt": utcnow_iso(),
        "createdBy": current_user["id"],
        "updatedAt": None,
        "generationMethod": "ai",
        "originalPrompt": body.prompt,
    }

    atomic_write_json(template_dir / TEMPLATE_META_FILENAME, meta)

    submissions_dir = _get_submissions_dir_for_tenant(request)
    return _load_template_with_interview(template_dir, submissions_dir)

@router.post("/{template_id}/regenerate")
def regenerate_template(
    template_id: str,
    body: RegenerateRequest,
    request: Request,
    current_user: dict = Depends(verify_tenant_match),
):
    if current_user["role"] not in ("admin", "superadmin"):
        raise HTTPException(status_code=403, detail="Not permitted")
    templates_dir = get_templates_dir(request)
    template_dir = templates_dir / template_id
    meta_path = template_dir / TEMPLATE_META_FILENAME
    if not meta_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")

    api_key = os.environ.get("OPENAI_API_KEY", settings.OPENAI_API_KEY)
    if not api_key:
        raise HTTPException(status_code=501, detail="AI generation not configured")

    model = os.environ.get("OPENAI_MODEL", "gpt-4o")
    meta = json.loads(meta_path.read_text())

    ai_result = _call_ai(body.prompt, model)

    document_content = ai_result.get("document_content", "")
    raw_questions = ai_result.get("questions", [])

    if not document_content:
        raise HTTPException(status_code=500, detail="AI did not generate document content")
    if not raw_questions:
        raise HTTPException(status_code=500, detail="AI did not generate interview questions")

    try:
        fields = validate_questions(raw_questions)
    except ValueError as e:
        raise HTTPException(status_code=500, detail=f"AI generated invalid questions: {e}")

    upload_path = template_dir / TEMPLATE_DOCX_FILENAME
    try:
        _create_docx_from_content(document_content, upload_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create document: {e}")

    interview_path = template_dir / TEMPLATE_INTERVIEW_FILENAME
    if interview_path.exists():
        interview = json.loads(interview_path.read_text())
    else:
        interview = {
            "schemaVersion": 1,
            "id": f"{template_id}_interview",
            "version": 1,
        }
    interview["components"] = fields
    # Auto-increment interview version (#4)
    interview["version"] = interview.get("version", 0) + 1
    # Remove title/description from interview (#6)
    interview.pop("title", None)
    interview.pop("description", None)
    interview.pop("$schema", None)
    atomic_write_json(interview_path, interview)

    meta["originalPrompt"] = body.prompt
    meta["generationMethod"] = "ai"
    meta["updatedAt"] = utcnow_iso()

    atomic_write_json(meta_path, meta)

    submissions_dir = _get_submissions_dir_for_tenant(request)
    return _load_template_with_interview(template_dir, submissions_dir)

# Serve files from a template subdirectory safely
@router.get("/download/{template_id}/{filename}")
def download_generated_file(
    template_id: str,
    filename: str,
    request: Request,
    current_user: dict = Depends(verify_tenant_match),
):
    if not re.match(r"^[\w\-. ]+\.(json|docx)$", filename):
        raise HTTPException(status_code=400, detail="Invalid filename")
    templates_dir = get_templates_dir(request)
    template_dir = templates_dir / template_id
    file_path = os.path.abspath(os.path.join(str(template_dir), filename))
    if not file_path.startswith(str(template_dir.resolve()) + os.sep):
        raise HTTPException(status_code=400, detail="Invalid path")
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path, filename=filename, media_type="application/octet-stream")
